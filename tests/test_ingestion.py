"""Ingestion: every format, and every way a file can disappoint us."""

from __future__ import annotations

import pytest

from src.ingestion.loader import classify_role, load_bundle, load_document
from src.ingestion.spreadsheet_parser import normalise_headers, parse_spreadsheet
from src.ingestion.types import UnreadableFile, UnsupportedFile
from src.models.evidence import SourceType
from tests.factories import write_pdf, write_pptx


def test_pdf_pages_keep_their_numbers(deck_path):
    document = load_document(deck_path, SourceType.PITCH_DECK)
    assert len(document.segments) == 4
    assert document.segments[3].index == 4
    assert "Raising $6M Series A" in document.segments[3].text
    hits = document.find("Raising $6M")
    assert hits and hits[0][0].index == 4


def test_pptx_extracts_text_tables_and_notes(tmp_path):
    path = write_pptx(
        tmp_path / "deck.pptx",
        [
            ("Testco", ["Rapid sepsis detection"], "Speaker note: mention the pilots"),
            ("The ask", ["Raising $4M seed on a SAFE with a $15M cap"], ""),
        ],
    )
    document = load_document(path, SourceType.PITCH_DECK)
    assert len(document.segments) == 2
    assert document.segments[0].title == "Testco"
    assert "Speaker note" in document.segments[0].notes
    assert "speaker notes" in document.segments[0].full_text()
    assert "$4M" in document.segments[1].full_text()


def test_image_only_pdf_is_flagged_not_dropped(tmp_path):
    path = write_pdf(tmp_path / "sparse.pdf", [["Testco"], [""], ["The ask", "Raising $3M"]])
    document = load_document(path, SourceType.PITCH_DECK)
    assert any("image-only" in w.message for w in document.warnings)
    # The readable pages survive.
    assert "Raising $3M" in document.text


def test_password_protected_pdf_explains_itself(tmp_path):
    import pypdf

    plain = write_pdf(tmp_path / "plain.pdf", [["Secret deck"]])
    reader = pypdf.PdfReader(str(plain))
    writer = pypdf.PdfWriter()
    for page in reader.pages:
        writer.add_page(page)
    writer.encrypt("hunter2")
    locked = tmp_path / "locked.pdf"
    with locked.open("wb") as handle:
        writer.write(handle)

    with pytest.raises(UnreadableFile) as excinfo:
        load_document(locked, SourceType.PITCH_DECK)
    assert "password" in str(excinfo.value).lower()


def test_corrupt_pdf_raises_a_readable_error(tmp_path):
    path = tmp_path / "broken.pdf"
    path.write_bytes(b"%PDF-1.4 this is not really a pdf")
    with pytest.raises(UnreadableFile):
        load_document(path, SourceType.PITCH_DECK)


def test_unsupported_extension_is_rejected(tmp_path):
    path = tmp_path / "notes.rtf"
    path.write_text("hello", encoding="utf-8")
    with pytest.raises(UnsupportedFile):
        load_document(path)


def test_empty_file_is_rejected(tmp_path):
    path = tmp_path / "empty.csv"
    path.write_bytes(b"")
    with pytest.raises(UnreadableFile):
        load_document(path)


def test_csv_headers_are_normalised(investors_csv):
    document = parse_spreadsheet(investors_csv, SourceType.INVESTOR_LIST)
    assert len(document.rows) == 7
    first = document.rows[0]
    assert first.get("investor_name") == "Bigname Global Partners"
    assert first.get("check_size") == "$5M-$20M"
    assert first.get("stage_focus") == "Series B, Series C"
    # The original heading is preserved alongside the canonical key.
    assert first.values["Investor"] == "Bigname Global Partners"


def test_header_synonyms_map_to_canonical_keys():
    mapping = normalise_headers(["Firm Name", "Typical Ticket", "Round Stage", "Pipeline Stage"])
    assert mapping[1] == "check_size"
    assert mapping[2] == "stage_focus"
    assert mapping[3] == "status"


def test_spreadsheet_with_preamble_rows_finds_the_header(tmp_path):
    path = tmp_path / "messy.csv"
    path.write_text(
        "Investor pipeline export\n"
        "Generated 2026-08-01\n"
        "Investor,Check Size,Stage\n"
        "Northlight Diagnostics Fund,$2M-$4M,Series A\n",
        encoding="utf-8",
    )
    document = parse_spreadsheet(path, SourceType.INVESTOR_LIST)
    assert len(document.rows) == 1
    assert document.rows[0].get("investor_name") == "Northlight Diagnostics Fund"


def test_xlsx_is_read(tmp_path):
    from openpyxl import Workbook

    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Pipeline"
    sheet.append(["Investor", "Check Size", "Leads Rounds"])
    sheet.append(["Northlight Diagnostics Fund", "$2M-$4M", "Yes"])
    path = tmp_path / "pipeline.xlsx"
    workbook.save(str(path))

    document = load_document(path)
    assert document.rows[0].get("investor_name") == "Northlight Diagnostics Fund"
    assert document.segments[0].title == "Pipeline"


def test_docx_sections_are_preserved(tmp_path):
    import docx

    document = docx.Document()
    document.add_heading("Northlight Diagnostics Fund", level=2)
    document.add_paragraph("Led the Vessl Dx Series A in 2025.")
    path = tmp_path / "research.docx"
    document.save(str(path))

    parsed = load_document(path)
    assert parsed.segments[0].title == "Northlight Diagnostics Fund"
    assert "Vessl Dx" in parsed.text


def test_role_classification_uses_the_filename(tmp_path):
    assert classify_role(tmp_path / "crm_export.csv") == SourceType.CRM_EXPORT
    assert classify_role(tmp_path / "investor_targets.csv") == SourceType.INVESTOR_LIST
    assert classify_role(tmp_path / "meeting_notes.md") == SourceType.MEETING_NOTES
    # A spreadsheet name that hints at notes is still read as a list.
    assert classify_role(tmp_path / "pipeline.xlsx") == SourceType.INVESTOR_LIST


def test_one_bad_file_does_not_stop_the_bundle(tmp_path, deck_path, investors_csv):
    broken = tmp_path / "broken.docx"
    broken.write_bytes(b"not a docx")
    bundle = load_bundle(deck_path, [investors_csv, broken])

    assert bundle.deck is not None
    assert len(bundle.supporting) == 1
    assert bundle.failed and "broken.docx" in bundle.failed[0][0]
    assert any("Skipped broken.docx" in w.message for w in bundle.warnings)


def test_missing_deck_is_reported_not_raised(tmp_path, investors_csv):
    bundle = load_bundle(tmp_path / "nope.pdf", [investors_csv])
    assert bundle.deck is None
    assert any(w.severity == "error" for w in bundle.warnings)
    assert len(bundle.supporting) == 1
