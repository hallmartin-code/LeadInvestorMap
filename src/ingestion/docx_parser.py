"""Word ingestion for investor notes, research memos and diligence documents.

Headings become section boundaries so a claim can cite "the Redwood section" rather than
just "the notes file".
"""

from __future__ import annotations

from pathlib import Path

from ..models.evidence import SourceType, Warning
from ..utils.text import clean
from .types import ParsedDocument, Segment, UnreadableFile


def parse_docx(path: str | Path, source_type: SourceType = SourceType.INVESTOR_RESEARCH) -> ParsedDocument:
    path = Path(path)
    try:
        import docx
    except ImportError as exc:  # pragma: no cover - dependency is declared
        raise UnreadableFile("python-docx is not installed.") from exc

    try:
        document = docx.Document(str(path))
    except Exception as exc:
        raise UnreadableFile(f"{path.name} could not be opened as a Word document: {exc}") from exc

    doc = ParsedDocument(path=path, name=path.name, source_type=source_type)

    sections: list[Segment] = []
    current = Segment(index=1, kind="section", title="")
    buffer: list[str] = []

    for paragraph in document.paragraphs:
        text = clean(paragraph.text)
        style = (paragraph.style.name or "").lower() if paragraph.style is not None else ""
        if not text:
            continue
        if style.startswith("heading") or style == "title":
            if buffer:
                current.text = "\n".join(buffer)
                sections.append(current)
                buffer = []
                current = Segment(index=len(sections) + 1, kind="section", title=text)
            else:
                current.title = text
        else:
            buffer.append(text)

    if buffer or current.title:
        current.text = "\n".join(buffer)
        sections.append(current)

    for table in document.tables:
        rows = [[clean(cell.text) for cell in row.cells] for row in table.rows]
        if any(any(cell for cell in row) for row in rows):
            if sections:
                sections[-1].tables.append(rows)
            else:
                sections.append(Segment(index=1, kind="section", tables=[rows]))

    doc.segments = sections or [Segment(index=1, kind="section", text="")]
    if doc.total_chars == 0:
        doc.warnings.append(
            Warning(
                severity="warning",
                stage="ingestion",
                message=f"No readable text was found in {path.name}.",
            )
        )
    doc.metadata = {"sections": len(doc.segments)}
    return doc
